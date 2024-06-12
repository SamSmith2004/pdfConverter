import fitz
import os
from docx import Document
import io
from PIL import Image
"""
Guide: put pdfs in pdfIn folder and run the script
"""

# Get filename from user
fileName = input('What is the filename(without .pdf)?\n')

# Open the PDF file
doc_pdf = fitz.open('./pdfIn/' + fileName + '.pdf')

# Create a new DOC file
doc = Document()

# For each page in the PDF
for i in range(len(doc_pdf)):
    page = doc_pdf.load_page(i)

    # Extract the images
    image_list = page.get_images(full=True)
    for img_index, img in enumerate(image_list):
        xref = img[0]
        base = img[1]
        img_data = doc_pdf.extract_image(xref)
        img_data = img_data["image"]

        # Save the image data to an io.BytesIO object
        image_stream = io.BytesIO(img_data)
        image = Image.open(image_stream)
        image_file_path = f"temp_image_{i}_{img_index}.png"
        image.save(image_file_path)

        # Add the image to the DOC and delete the temporary image file
        doc.add_picture(image_file_path)
        os.remove(image_file_path)

    # Extract the text
    text = page.get_text()

    # Insert the text into the DOC file
    doc.add_paragraph(text)

# Save the DOC file
doc.save('./docOut/' + fileName + '.docx')